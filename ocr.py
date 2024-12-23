import streamlit as st
import ollama
from PIL import Image
import io
import time
from pdf2image import convert_from_bytes
import mimetypes
from docx import Document
import unicodedata


# Remove surrogate characters from text
def remove_surrogates(text):
    return ''.join(char for char in text if not unicodedata.category(char).startswith("Cs"))


# Page configuration
st.set_page_config(
    page_title="AI-Powered Text Extraction Tool",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)


# Enhanced Debugging Function
def log_timing(step_name, start_time):
    elapsed_time = time.time() - start_time
    st.write(f"⏱️ {step_name} took {elapsed_time:.2f} seconds")


# Clear data on file upload
def clear_results():
    st.session_state["ocr_result"] = None


# Supported file formats
SUPPORTED_FORMATS = ["jpg", "jpeg", "png", "gif", "jfif", "heic", "pdf", "tiff"]


# File validation function
def validate_file(uploaded_file):
    mime_type, _ = mimetypes.guess_type(uploaded_file.name)
    supported_mime_types = {
        "image/jpeg": "jpeg",
        "image/png": "png",
        "image/gif": "gif",
        "image/tiff": "tiff",
        "application/pdf": "pdf",
    }
    if mime_type in supported_mime_types:
        return supported_mime_types[mime_type]
    return None


# Process uploaded file
def process_uploaded_file(uploaded_file, file_extension):
    uploaded_file.seek(0)  # Reset the file pointer
    if file_extension in ["jpg", "jpeg", "png", "gif", "jfif", "heic", "tiff"]:
        return Image.open(uploaded_file)
    elif file_extension == "pdf":
        try:
            pdf_bytes = uploaded_file.read()
            images = convert_from_bytes(pdf_bytes, dpi=300)
            return images[0]  # Process the first page only
        except Exception as e:
            raise ValueError(f"Failed to process PDF file. Error: {e}")
    else:
        raise ValueError("Unsupported file format")


# Custom CSS Styling
st.markdown(
    """
    <style>
    body {
        background: linear-gradient(135deg, #6a11cb 0%, #2575fc 100%);
        font-family: 'Roboto', sans-serif;
        color: #fff;
    }
    h1 {
        color: #ffffff;
        text-align: center;
        font-size: 3em;
        margin-bottom: 0.3em;
        font-weight: bold;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
    }
    .description {
        text-align: center;
        font-size: 1.3em;
        color: #e0e0e0;
        margin-bottom: 1.5em;
    }
    .ocr-result {
        font-family: 'Roboto', sans-serif;
        font-size: 1.1em;
        line-height: 1.6;
        color: #000;
        background: #ffffff;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# App title and description
start_time = time.time()
st.markdown(
    """
    <h1>AI-Powered Text Extraction Tool</h1>
    <p class="description">Extract text from images and PDFs with ease, powered by cutting-edge AI.</p>
    """,
    unsafe_allow_html=True,
)
log_timing("Page Setup", start_time)


# Sidebar for uploading files
with st.sidebar:
    st.markdown("<div class='sidebar-header'>Upload Your File</div>", unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Choose a file...", type=SUPPORTED_FORMATS, on_change=clear_results)
    log_timing("Sidebar Rendering", start_time)

    if uploaded_file:
        try:
            # Determine file extension
            file_extension = validate_file(uploaded_file)
            if not file_extension:
                st.error("Unsupported file format. Please upload a valid file.")
                st.stop()

            # Process file into an image
            image = process_uploaded_file(uploaded_file, file_extension)
            st.image(image, caption="Uploaded File Preview", width=300)

            if st.button("Extract Text 🔍"):
                with st.spinner("Processing your file..."):
                    try:
                        # Convert image to bytes in a valid format for the API
                        img_bytes = io.BytesIO()
                        image.save(img_bytes, format="PNG")  # Ensure the image is encoded as PNG
                        img_bytes.seek(0)  # Reset the pointer for reading

                        # Extract text from the image
                        response = ollama.chat(
                            model="llama3.2-vision",
                            messages=[
                                {
                                    "role": "user",
                                    "content": (
                                        "Analyze the text in the provided image. Extract all readable content "
                                        "and present it in a structured Markdown format."
                                    ),
                                    "images": [img_bytes.read()],  # Pass valid image bytes for processing
                                }
                            ],
                        )

                        # Sanitize the response to remove invalid characters
                        sanitized_response = remove_surrogates(response.message.content)
                        st.session_state["ocr_result"] = sanitized_response

                    except Exception as e:
                        st.error(f"Error extracting text: {e}")
        except Exception as e:
            st.error(f"Error processing file: {e}")


# Display OCR result and download options
if "ocr_result" in st.session_state and st.session_state["ocr_result"]:
    sanitized_result = remove_surrogates(st.session_state["ocr_result"])
    st.markdown(f"<div class='ocr-result'>{sanitized_result}</div>", unsafe_allow_html=True)

    # Download as .txt
    result_bytes = io.BytesIO(sanitized_result.encode("utf-8", errors="replace"))
    st.download_button(
        label="Download Results as .txt",
        data=result_bytes,
        file_name="ocr_result.txt",
        mime="text/plain",
    )

    # Download as .docx
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    doc.add_paragraph(sanitized_result)
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    st.download_button(
        label="Download Results as .docx",
        data=docx_bytes,
        file_name="ocr_result.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
else:
    st.info("Upload a file and click 'Extract Text' to see results.")
